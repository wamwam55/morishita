# -*- coding: utf-8 -*-
"""マネーフォワード クラウド会計 かんたん操作マニュアル — 組み立て

  /tmp/_docvenv/bin/python assemble.py
  → ~/Downloads/マネーフォワード会計_かんたん操作マニュアル_2026-08.docx
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from common import setup_styles, page_setup, add_page_number_footer, DOC_TITLE

import build_intro
import build_ch1
import build_ch2to6
import build_accounts
import build_ch8to10
import build_appendix


def main():
    doc = Document()
    page_setup(doc)
    setup_styles(doc)
    add_page_number_footer(doc)

    build_intro.build(doc)      # 表紙 / このマニュアルについて / 目次 / 第0章
    build_ch1.build(doc)        # 第1章 初期設定
    build_ch2to6.build(doc)     # 第2〜6章
    build_accounts.build(doc)   # 第7章 勘定科目
    build_ch8to10.build(doc)    # 第8〜10章
    build_appendix.build(doc)   # 付録A〜D

    core = doc.core_properties
    core.title = DOC_TITLE
    core.author = "森下知幸税理士・社労士事務所"
    core.comments = "2026年8月版（第1.0版）"

    out = os.path.expanduser("~/Downloads/マネーフォワード会計_かんたん操作マニュアル_2026-08.docx")
    doc.save(out)

    tables = len(doc.tables)
    paras = len(doc.paragraphs)
    print(f"saved: {out}")
    print(f"tables={tables} paragraphs={paras} size={os.path.getsize(out):,} bytes")


if __name__ == "__main__":
    main()
