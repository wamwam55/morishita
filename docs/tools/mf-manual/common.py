# -*- coding: utf-8 -*-
"""マネーフォワード クラウド会計 初心者向け操作マニュアル — 共通部品

デザインは森下知幸税理士・社労士事務所サイト（morishita-tax.jp）の配色に合わせている。
  ティール  #4A7C7E / #5D9B9B  … 見出し・表ヘッダ・ポイント
  トープ    #8B7962            … 注意ボックス
  見出し黒  #1A1A1A            … 本文見出し
  罫線      #E5E5E5 / 背景 #FAFAFA
"""
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

JP_FONT = "游ゴシック"

# --- サイト配色 ---
TEAL_HEX = "4A7C7E"
TEAL_LIGHT_HEX = "5D9B9B"
TAUPE_HEX = "8B7962"
INK_HEX = "1A1A1A"
GRAY_HEX = "666666"

TEAL = RGBColor(0x4A, 0x7C, 0x7E)
TEAL_LIGHT = RGBColor(0x5D, 0x9B, 0x9B)
TAUPE = RGBColor(0x8B, 0x79, 0x62)
INK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)

BOX_TEAL_FILL = "EDF4F4"
BOX_TAUPE_FILL = "F7F3EE"
BAND_FILL = "FAFAFA"

DOC_TITLE = "マネーフォワード クラウド会計 かんたん操作マニュアル"


def set_jp(run, size=None, bold=None, color=None, font=JP_FONT):
    run.font.name = font
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font)
    r.set(qn('w:ascii'), font)
    r.set(qn('w:hAnsi'), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def borders(paragraph, color=TEAL_HEX, sz=6, sides=("left",)):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for s in sides:
        e = OxmlElement(f'w:{s}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '4')
        e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)


def spacing(paragraph, before=0, after=6, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def setup_styles(doc):
    st = doc.styles['Normal']
    st.font.name = JP_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), JP_FONT)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(4)

    for name, size, color in (("Heading 1", 17, TEAL),
                              ("Heading 2", 13.5, TEAL),
                              ("Heading 3", 11.5, INK)):
        s = doc.styles[name]
        s.font.name = JP_FONT
        s.element.rPr.rFonts.set(qn('w:eastAsia'), JP_FONT)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True


def page_setup(doc):
    for s in doc.sections:
        s.page_width = Mm(210)
        s.page_height = Mm(297)
        s.top_margin = Mm(20)
        s.bottom_margin = Mm(18)
        s.left_margin = Mm(20)
        s.right_margin = Mm(20)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_jp(run, 9)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    p._p.append(fld)
    r2 = p.add_run("　/　" + DOC_TITLE)
    set_jp(r2, 8, color=GRAY)


# ---------- 段落ヘルパ ----------
def h1(doc, text, new_page=True):
    """章見出し。new_page=True で必ず改ページしてから始める。

    空の段落による改ページだと、直前の本文がページ末で終わったときに
    まるごと白紙のページができてしまう。page_break_before なら発生しない。
    """
    p = doc.add_heading(level=1)
    borders(p, color=TEAL_HEX, sz=18, sides=("left",))
    p.paragraph_format.left_indent = Mm(3)
    if new_page:
        p.paragraph_format.page_break_before = True
    set_jp(p.add_run(text), 17, True, TEAL)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    set_jp(p.add_run(text), 13.5, True, TEAL)
    return p


def h3(doc, text):
    p = doc.add_heading(level=3)
    set_jp(p.add_run(text), 11.5, True, INK)
    return p


def para(doc, text, size=10.5, bold=False, align=None, after=4, color=None):
    p = doc.add_paragraph()
    set_jp(p.add_run(text), size, bold, color)
    if align:
        p.alignment = align
    spacing(p, after=after)
    return p


def steps(doc, items):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Mm(8)
        pf.first_line_indent = Mm(-8)
        spacing(p, after=3)
        set_jp(p.add_run(f"{i}. "), 10.5, True, TEAL)
        set_jp(p.add_run(t), 10.5)


def bullets(doc, items, indent=6, keep=False):
    """keep=True にすると箇条書きがページをまたいで分断されない。"""
    for i, t in enumerate(items):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Mm(indent + 4)
        pf.first_line_indent = Mm(-4)
        if keep and i < len(items) - 1:
            pf.keep_with_next = True
        spacing(p, after=2)
        set_jp(p.add_run("・"), 10.5, color=TEAL)
        set_jp(p.add_run(t), 10.5)


def box(doc, label, text, fill=BOX_TEAL_FILL, color=TEAL_HEX):
    p = doc.add_paragraph()
    shade(p, fill)
    borders(p, color=color, sz=18, sides=("left",))
    pf = p.paragraph_format
    pf.left_indent = Mm(3)
    pf.right_indent = Mm(2)
    spacing(p, before=6, after=8)
    set_jp(p.add_run(f"{label}　"), 10.5, True, RGBColor.from_string(color))
    set_jp(p.add_run(text), 10.5)
    return p


def point(doc, text):
    return box(doc, "◆ ポイント", text, BOX_TEAL_FILL, TEAL_HEX)


def warn(doc, text):
    return box(doc, "▲ 注意", text, BOX_TAUPE_FILL, TAUPE_HEX)


def table(doc, headers, rows, widths=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        set_jp(p.add_run(htxt), size, True, RGBColor(0xFF, 0xFF, 0xFF))
        spacing(p, after=2)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), TEAL_HEX)
        tcPr.append(shd)
    for ri, r in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_jp(p.add_run(str(val)), size)
            spacing(p, after=2)
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'F4F8F8')
                tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
