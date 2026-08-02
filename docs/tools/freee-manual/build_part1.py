# -*- coding: utf-8 -*-
"""freee会計 初心者向け操作マニュアル — 共通部品と表紙〜第1章"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

JP_FONT = "游ゴシック"
ACCENT = RGBColor(0x1F, 0x6F, 0x8B)
DARK = RGBColor(0x22, 0x22, 0x22)


def set_jp(run, size=None, bold=None, color=None, font=JP_FONT):
    run.font.name = font
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font)
    r.set(qn('w:ascii'), font)
    r.set(qn('w:hAnsi'), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def borders(paragraph, color="1F6F8B", sz=6, sides=("left",)):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for s in sides:
        e = OxmlElement(f'w:{s}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '4')
        e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)


def spacing(paragraph, before=0, after=6, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def setup_styles(doc):
    st = doc.styles['Normal']
    st.font.name = JP_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), JP_FONT)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(4)

    for name, size, color in (("Heading 1", 17, ACCENT),
                              ("Heading 2", 13.5, ACCENT),
                              ("Heading 3", 11.5, DARK)):
        s = doc.styles[name]
        s.font.name = JP_FONT
        s.element.rPr.rFonts.set(qn('w:eastAsia'), JP_FONT)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True


def page_setup(doc):
    for s in doc.sections:
        s.page_width = Mm(210)
        s.page_height = Mm(297)
        s.top_margin = Mm(20)
        s.bottom_margin = Mm(18)
        s.left_margin = Mm(20)
        s.right_margin = Mm(20)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_jp(run, 9)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    p._p.append(fld)
    r2 = p.add_run("　/　freee会計 かんたん操作マニュアル")
    set_jp(r2, 8, color=RGBColor(0x88, 0x88, 0x88))


# ---------- 段落ヘルパ ----------
def h1(doc, text):
    p = doc.add_heading(level=1)
    set_jp(p.add_run(text), 17, True, ACCENT)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    set_jp(p.add_run(text), 13.5, True, ACCENT)
    return p


def h3(doc, text):
    p = doc.add_heading(level=3)
    set_jp(p.add_run(text), 11.5, True, DARK)
    return p


def para(doc, text, size=10.5, bold=False, align=None, after=4):
    p = doc.add_paragraph()
    set_jp(p.add_run(text), size, bold)
    if align:
        p.alignment = align
    spacing(p, after=after)
    return p


def steps(doc, items):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Mm(8)
        pf.first_line_indent = Mm(-8)
        spacing(p, after=3)
        set_jp(p.add_run(f"{i}. "), 10.5, True, ACCENT)
        set_jp(p.add_run(t), 10.5)


def bullets(doc, items, indent=6):
    for t in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Mm(indent + 4)
        pf.first_line_indent = Mm(-4)
        spacing(p, after=2)
        set_jp(p.add_run("・"), 10.5, color=ACCENT)
        set_jp(p.add_run(t), 10.5)


def box(doc, label, text, fill="EAF3F7", color="1F6F8B"):
    p = doc.add_paragraph()
    shade(p, fill)
    borders(p, color=color, sz=18, sides=("left",))
    pf = p.paragraph_format
    pf.left_indent = Mm(3)
    pf.right_indent = Mm(2)
    spacing(p, before=6, after=8)
    set_jp(p.add_run(f"{label}　"), 10.5, True, RGBColor.from_string(color))
    set_jp(p.add_run(text), 10.5)
    return p


def point(doc, text):
    return box(doc, "◆ ポイント", text, "EAF3F7", "1F6F8B")


def warn(doc, text):
    return box(doc, "▲ 注意", text, "FDF1E7", "C0651A")


def table(doc, headers, rows, widths=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        set_jp(p.add_run(htxt), size, True, RGBColor(0xFF, 0xFF, 0xFF))
        spacing(p, after=2)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '1F6F8B')
        tcPr.append(shd)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_jp(p.add_run(str(val)), size)
            spacing(p, after=2)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------- 本体 ----------
def build_cover_and_ch1(doc):
    # 表紙
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_jp(p.add_run("freee会計"), 34, True, ACCENT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_jp(p.add_run("かんたん操作マニュアル"), 26, True, ACCENT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=6, after=30)
    set_jp(p.add_run("〜 はじめて使う方のための手引き 〜"), 13, False, DARK)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade(p, "EAF3F7")
    spacing(p, before=10, after=10)
    set_jp(p.add_run("この1冊で「つなぐ → 決める → 見る」ができるようになります"), 11, True, ACCENT)

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_jp(p.add_run("森下知幸税理士・社労士事務所"), 14, True, DARK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_jp(p.add_run("2026年8月版（第1.0版）"), 10.5, False, DARK)
    pagebreak(doc)

    # このマニュアルについて
    h1(doc, "このマニュアルについて")
    para(doc, "freee会計をはじめて使う方が、ひとりで日々の経理を回せるようになることを目的にまとめています。"
              "むずかしい簿記の知識がなくても読み進められるように、画面の名前とボタンの名前をそのまま書いています。")
    h3(doc, "対象となる方")
    bullets(doc, ["これから freee会計 を使いはじめる方",
                  "使いはじめたばかりで、どこを押せばよいか迷っている方",
                  "スマホでレシートを処理したい方"])
    h3(doc, "読み方のコツ")
    bullets(doc, ["最初から順番に読む必要はありません。必要な章だけ開いてください。",
                  "［　］で囲んだ言葉は、freeeの画面上のメニュー名・ボタン名です。そのまま画面で探してください。",
                  "「◆ ポイント」は覚えておくと楽になること、「▲ 注意」は間違えやすいところです。"])
    warn(doc, "本書の画面名は2026年8月時点の freee会計（新しいメニューバー）に合わせています。"
              "freeeは頻繁に改良されるため、名称が変わることがあります。その場合は［　］内の言葉を手がかりに"
              "画面上部の検索から探してください。金額・税務の判断は、最終的に必ず顧問税理士へご確認ください。")

    h2(doc, "目次")
    toc = [
        ("第0章", "はじめに — freee会計の考え方（3分）"),
        ("第1章", "データ連携の方法（銀行・カードをつなぐ）"),
        ("第2章", "自動入力の操作方法（自動で経理）"),
        ("第3章", "入力ルールの作成方法（自動登録ルール）"),
        ("第4章", "会計データの確認方法"),
        ("第5章", "スマホでレシートを撮影して処理する方法"),
        ("第6章", "勘定科目の簡単な説明一覧"),
        ("第7章", "取引の消込（債権）の方法"),
        ("第8章", "口座振替の説明"),
        ("付録A", "月次の締め作業チェックリスト"),
        ("付録B", "用語集"),
        ("付録C", "よくある質問とつまずきポイント"),
        ("付録D", "困ったときの調べ方"),
    ]
    for no, title in toc:
        p = doc.add_paragraph()
        spacing(p, after=3)
        pf = p.paragraph_format
        pf.left_indent = Mm(4)
        set_jp(p.add_run(f"{no}　"), 10.5, True, ACCENT)
        set_jp(p.add_run(title), 10.5)
    pagebreak(doc)

    # 第0章
    h1(doc, "第0章　はじめに — freee会計の考え方")
    para(doc, "freee会計は「仕訳を手で打ち込むソフト」ではありません。"
              "銀行やクレジットカードの明細を自動で取り込み、そこから帳簿を作るソフトです。"
              "ここを understand しておくと、以降の操作がすべてつながって見えてきます。".replace("understand ", "理解"))

    h2(doc, "0-1　基本は「つなぐ → 決める → 見る」の3ステップ")
    table(doc,
          ["ステップ", "やること", "本書の章"],
          [["① つなぐ", "銀行口座・クレジットカードを freee につなぐ（最初の1回だけ）", "第1章"],
           ["② 決める", "取り込まれた明細に勘定科目をつけて［登録］する（毎日〜毎週）", "第2章・第3章"],
           ["③ 見る", "試算表などで数字を確認する（毎月）", "第4章"]],
          widths=[28, 118, 24])
    point(doc, "②の作業を毎回ゼロから考えなくて済むようにする仕組みが「自動登録ルール」（第3章）です。"
               "最初の1か月でルールを育てると、あとの作業時間が一気に減ります。")

    h2(doc, "0-2　最初に覚える3つの言葉")
    table(doc,
          ["言葉", "意味", "たとえると"],
          [["同期（どうき）", "銀行・カードのデータを freee に取り込むこと", "通帳を freee に見せる"],
           ["明細（めいさい）", "取り込まれた入出金1件ずつのデータ。まだ帳簿ではない", "通帳の1行"],
           ["取引（とりひき）", "明細に勘定科目をつけて帳簿に登録したもの", "帳簿に書き込んだ状態"]],
          widths=[30, 90, 50])
    para(doc, "freee では、明細を「取引」に変える作業＝日々の経理、と考えてください。")

    h2(doc, "0-3　メニューバーの全体像")
    para(doc, "2026年の freee会計 は、経理の流れに沿ってメニューが並んでいます。迷ったらこの表に戻ってください。")
    table(doc,
          ["メニュー", "ここでできること（主なもの）"],
          [["［取引入力］", "自動で経理／取引の一覧・登録／ファイルボックス／振替伝票"],
           ["［請求・入金］", "請求書の作成・発行／入金管理レポート（売掛金の管理）"],
           ["［発注・経費・支払］", "支払いの管理／支払管理レポート（買掛金の管理）"],
           ["［会計帳簿］", "貸借対照表／損益計算書／月次推移／仕訳帳／総勘定元帳／明細の一覧"],
           ["［分析・レポート］", "損益・収益・費用レポート／現預金レポート／資金繰りレポート"],
           ["［マスタ・口座］", "口座／勘定科目／税区分の設定"],
           ["［入力効率化］", "自動登録ルール／振替伝票テンプレート"],
           ["［その他設定］", "事業所の基本情報／事業所の詳細設定／開始残高／メンバー招待"]],
          widths=[42, 128])
    warn(doc, "以前の freee をご存じの方へ。かつての「設定」メニューや「レポート」メニューはなくなりました。"
              "口座と勘定科目は［マスタ・口座］、帳簿は［会計帳簿］、分析は［分析・レポート］に移っています。")

    h2(doc, "0-4　使いはじめる前に済ませておく初期設定")
    para(doc, "ここは顧問税理士と一緒に設定するのが確実です。すでに設定済みであれば読み飛ばしてください。")
    table(doc,
          ["順番", "設定すること", "場所"],
          [["1", "事業所の基本情報（事業所名・住所・業種など）", "［その他設定］→［事業所の基本情報］"],
           ["2", "会計期間（いつからいつまでを1年とするか）", "［その他設定］→［事業所の詳細設定］"],
           ["3", "消費税の設定（免税／簡易課税／一般課税、税込・税抜）", "［その他設定］→［事業所の詳細設定］"],
           ["4", "口座の登録（銀行・カード）", "［マスタ・口座］→［口座］（第1章）"],
           ["5", "開始残高（使いはじめ時点の資産・負債）", "［その他設定］→［開始残高］"],
           ["6", "勘定科目の確認・追加", "［マスタ・口座］→［勘定科目］（第6章）"]],
          widths=[14, 82, 74])
    bullets(doc, [
        "個人事業主の会計期間は1月1日〜12月31日で固定です（月日は変更できません）。",
        "消費税は「免税」「簡易課税」「一般課税」から選びます。freee の画面では『本則課税』ではなく『一般課税』と表示されます。",
        "免税事業者の方は「税込経理」を選んでください。",
        "開始残高が正しくないと、いくら日々の入力を頑張っても残高が合いません。ここは必ず税理士と確認してください。",
    ])
    pagebreak(doc)

    # 第1章
    h1(doc, "第1章　データ連携の方法")
    para(doc, "銀行口座とクレジットカードを freee につなぐと、明細が自動で入ってきます。"
              "最初の1回だけの作業ですが、ここが freee のいちばん大事な土台です。")

    h2(doc, "1-1　つなぐと何が変わるか")
    bullets(doc, ["通帳を見ながら1件ずつ手入力する作業がなくなる",
                  "入力もれ・金額の打ち間違いがなくなる",
                  "残高が自動で照合できるので、月末の確認が早く終わる",
                  "レシートとの突き合わせが楽になる"])

    h2(doc, "1-2　銀行口座をつなぐ")
    steps(doc, [
        "［マスタ・口座］メニュー →［口座］をクリックします。",
        "［＋登録］ボタンをクリックし、カテゴリを「銀行口座」、データの取得方法を「自動で取得」にして［登録］をクリックします。",
        "「連携先を選択してください」画面で、お使いの銀行名を検索してクリックします。",
        "「○○銀行と連携しますか？」画面で［次へ］をクリックします。",
        "銀行側の認証画面が開くので、インターネットバンキングのIDやパスワードなどを入力します。",
        "「○○銀行との接続を確認しています」と表示されるので、そのまま待ちます。",
        "「連携する銀行口座を選択してください」で、使う口座にチェックを入れて［次へ］をクリックします。",
        "「明細の取得を開始する日を設定してください」で開始日を選び［次へ］をクリックします。",
        "口座の一覧画面に戻り、明細の取り込みが始まります。",
    ])
    point(doc, "「API連携」に対応した銀行なら、freee にIDやパスワードを保存せずにデータを取得できます。"
               "銀行側で必要になる権限は「残高照会」と「入出金明細照会」の2つだけです。"
               "同じ銀行の別口座を足すときは、手順4で「登録済みのアカウントで連携」を選びます。")

    h2(doc, "1-3　クレジットカードをつなぐ")
    steps(doc, [
        "［マスタ・口座］メニュー →［口座］をクリックします。",
        "［＋登録］ボタンから、カテゴリを「クレジットカード」、取得方法を「自動で取得」にして［登録］をクリックします。",
        "「連携先を選択してください」でカード名を検索してクリックし、［次へ］をクリックします。",
        "カード会社のオンラインサービスのログイン情報（IDとパスワード）を入力して［次へ］をクリックします。"
        "（API連携対応のカードの場合は、カード会社の認証画面に移ります）",
        "接続確認の画面が出るので待ちます。追加の質問（秘密の質問など）が出たら入力して［更新］をクリックします。",
        "「連携するクレジットカードを選択してください」でカードにチェックを入れて［次へ］をクリックします。",
        "明細の取得開始日を選んで［次へ］をクリックします。",
        "口座の一覧画面に戻り、取り込みが始まります。",
    ])
    warn(doc, "カードの明細が freee に入ってくるのは、その明細が「確定」してからです。"
              "使った直後には出てきません。月末に「まだ出ていない」と慌てないでください。")

    h2(doc, "1-4　口座の一覧画面の見方")
    para(doc, "［マスタ・口座］→［口座］で開きます。ここが日々の健康診断の画面です。")
    table(doc,
          ["表示・ボタン", "意味・使い方"],
          [["同期失敗", "つながっていない口座の数。0になっているのが正常です"],
           ["未登録明細あり", "まだ帳簿にしていない明細がある口座の数。数字をクリックすると［自動で経理］へ飛べます"],
           ["残高ずれあり", "freeeの残高と実際の残高が違う口座の数。0を目指します（→ 4-7）"],
           ["登録残高", "freeeに登録した取引から計算した残高"],
           ["同期残高", "最後に同期したときの実際の口座残高（銀行口座のみ）"],
           ["［全口座同期］", "連携済みのすべての口座をまとめて同期します"],
           ["［同期］", "その口座だけを同期します"],
           ["［明細アップロード］", "CSVなどで明細を手動で取り込みます（→ 1-6）"],
           ["［＋フィルタ］", "状態・未登録明細・残高ずれなどで絞り込みます"]],
          widths=[40, 130])

    h2(doc, "1-5　うまくつながらないときは")
    h3(doc, "ケース1　そもそも同期が実行されていない")
    steps(doc, [
        "ホーム画面で該当口座の下にある「∨」のつまみをクリックし、［口座を同期］ボタンを押します。",
        "またはホーム画面左上の［全口座を同期］ボタンを押します。",
        "毎回自動で取り込みたい場合は、口座名をクリック →［口座設定］→「明細を取り込む方法を選びましょう」で"
        "「オンラインサービスと同期する」を選びます。",
    ])
    h3(doc, "ケース2　古い明細が取れない")
    para(doc, "銀行ごとに「さかのぼれる期間」が決まっており、それより前の明細は取り込めません。"
              "この場合は 1-6 の明細アップロードで補います。")
    h3(doc, "ケース3　カードの未確定明細がある")
    para(doc, "支払金額が確定するまで待ってから、もう一度同期してください。"
              "口座の詳細画面の上部にある［未確定明細対応可否］で、そのカードが未確定明細に対応しているか確認できます。")
    h3(doc, "エラーメッセージが出ている場合")
    steps(doc, [
        "ホーム画面上部に出ているエラーメッセージをクリックします。",
        "トラブルシューティング画面に、原因と解決のヒントが表示されます。",
        "画面上部の案内にしたがって対処します。多くはパスワード変更や追加認証が原因です。",
    ])
    h3(doc, "ログイン情報を入れ直す（再連携）")
    steps(doc, [
        "［マスタ・口座］→［口座］で対象の口座の行をクリックします。",
        "左上の［口座設定］ボタンをクリックします。",
        "「ログイン情報を入力しましょう」の［アカウント情報を変更する］をクリックします。",
        "API連携なら［認証ページへ］、通常連携なら正しいIDとパスワードを入力して［更新］をクリックします。",
        "［口座を保存する］ボタンをクリックします。",
    ])
    point(doc, "銀行側でパスワードを変更したら、freee 側も必ず入れ直してください。"
               "放置すると同期が止まり、その間の明細が抜けたままになります。")

    h2(doc, "1-6　CSVで明細を取り込む（同期できないとき）")
    steps(doc, [
        "ホーム画面左側の口座一覧から、取り込みたい口座の［∨］ボタンをクリックします。",
        "［明細のアップロード］ボタンをクリックします。",
        "［ファイルを選択］ボタンで、銀行からダウンロードした明細ファイルを選びます。",
        "ファイルの種類を選びます。分からないときは「ご自身で作成したCSV（新規のフォーマット）」を選びます。",
        "日付の形式・金額が入っている列・明細の並び順を指定します。",
        "プレビューを見ながら見出し（日付・金額など）を指定し、［明細を取り込む］ボタンをクリックします。",
        "取り込み結果を確認します。2回目以降は「前回と同じフォーマット」が使えます。",
    ])
    warn(doc, "「出金額」と「入金額」の列を逆に指定してしまう事故がいちばん多いです。取り込む前にプレビューで必ず確認してください。"
              "なお対応形式は CSV・OFX で、サイズは5MBまでです。Excelは保存時に拡張子を .csv にして変換してください。")
    point(doc, "手動でアップロードした明細には、明細行の右側に「手」のアイコンが付きます。"
               "このアイコンが付いた明細は、あとから削除・編集ができます（同期で入った明細は削除できません）。")

    h2(doc, "1-7　第1章のチェックリスト")
    bullets(doc, ["事業で使う銀行口座をすべて登録した",
                  "事業で使うクレジットカードをすべて登録した",
                  "口座一覧の「同期失敗」が0になっている",
                  "同期できない口座は、CSVで取り込む段取りを決めた"])
    pagebreak(doc)


if __name__ == "__main__":
    doc = Document()
    setup_styles(doc)
    page_setup(doc)
    add_page_number_footer(doc)
    build_cover_and_ch1(doc)
    os.makedirs("/tmp/freee_manual", exist_ok=True)
    doc.save("/tmp/freee_manual/_part1.docx")
    print("part1 OK:", len(doc.paragraphs), "paragraphs")
