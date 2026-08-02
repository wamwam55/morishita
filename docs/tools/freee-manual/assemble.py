# -*- coding: utf-8 -*-
import os
from docx import Document
from build_part1 import setup_styles, page_setup, add_page_number_footer, build_cover_and_ch1
from build_part2 import build_ch2, build_ch3, build_ch4
from build_part3 import build_ch5, build_ch6
from build_part4 import build_ch7, build_ch8
from build_part5 import build_appendix_a, build_appendix_b, build_appendix_c, build_appendix_d

OUT = os.path.expanduser("~/Downloads/freee会計_かんたん操作マニュアル_2026-08.docx")

doc = Document()
setup_styles(doc)
page_setup(doc)
add_page_number_footer(doc)

core = doc.core_properties
core.title = "freee会計 かんたん操作マニュアル"
core.author = "森下知幸税理士・社労士事務所"
core.comments = "2026年8月版（第1.0版）"

build_cover_and_ch1(doc)
build_ch2(doc)
build_ch3(doc)
build_ch4(doc)
build_ch5(doc)
build_ch6(doc)
build_ch7(doc)
build_ch8(doc)
build_appendix_a(doc)
build_appendix_b(doc)
build_appendix_c(doc)
build_appendix_d(doc)

doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
